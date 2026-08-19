"""
JARVIS Word Document Tools
Handles .docx creation, editing, and manipulation
"""

from pathlib import Path
from typing import Optional, List, Dict, Any


def create_document(
    title: str,
    content: List[str],
    output_path: Optional[str] = None
) -> str:
    """Create a new Word document with title and content."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        
        doc = Document()
        doc.add_heading(title, 0)
        
        for line in content:
            if line.strip():
                doc.add_paragraph(line)
        
        if not output_path:
            output_path = f"{title.replace(' ', '_')}.docx"
        
        doc.save(output_path)
        return f"Document created: {output_path}"
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Error creating document: {e}"


def extract_text(docx_path: str) -> str:
    """Extract all text from a Word document."""
    try:
        from docx import Document
        
        doc = Document(docx_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        return "\n".join(text)
    except ImportError:
        return "Error: python-docx not installed"
    except Exception as e:
        return f"Error extracting text: {e}"


def add_table(docx_path: str, rows: List[List[str]], cols: int) -> str:
    """Add a table to an existing document."""
    try:
        from docx import Document
        
        doc = Document(docx_path)
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Light Grid Accent 1'
        
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                table.rows[i].cells[j].text = str(cell)
        
        doc.save(docx_path)
        return f"Table added to {docx_path}"
    except Exception as e:
        return f"Error adding table: {e}"
